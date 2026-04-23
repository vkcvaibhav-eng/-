import streamlit as st
import google.generativeai as genai
import sqlite3
import datetime
from PIL import Image
import io
import os
import PyPDF2
from docx import Document as DocxReader
import urllib.request

# New Imports for DOCX Generation
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# Database Setup for Archiving
# ==========================================
DB_FILE = "sadar_nondh_archive.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS archive 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  date TEXT, 
                  month TEXT, 
                  year TEXT, 
                  subject TEXT, 
                  content TEXT)''')
    conn.commit()
    conn.close()

def save_to_db(subject, content):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    now = datetime.datetime.now()
    c.execute("INSERT INTO archive (date, month, year, subject, content) VALUES (?, ?, ?, ?, ?)", 
              (now.strftime("%d/%m/%Y"), now.strftime("%m"), now.strftime("%Y"), subject, content))
    conn.commit()
    conn.close()

def get_archives(month, year):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    if month == "All":
        c.execute("SELECT date, subject, content FROM archive WHERE year=?", (year,))
    else:
        c.execute("SELECT date, subject, content FROM archive WHERE month=? AND year=?", (month, year))
    data = c.fetchall()
    conn.close()
    return data

init_db()

# ==========================================
# Permanent Attachments Extraction
# ==========================================
@st.cache_data
def load_permanent_context():
    statute_text = "Statute 121 Rules:\n"
    sample_text = "Sample Nondh Format:\n"
    
    if os.path.exists("121_Statutes_uploaded.pdf"):
        try:
            with open("121_Statutes_uploaded.pdf", "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    statute_text += page.extract_text() + "\n"
        except Exception as e:
            st.warning(f"Could not load 121 Statutes PDF: {e}")

    if os.path.exists("sample_nondh_uploaded.docx"):
        try:
            doc = DocxReader("sample_nondh_uploaded.docx")
            for para in doc.paragraphs:
                sample_text += para.text + "\n"
        except Exception as e:
            st.warning(f"Could not load sample_nondh DOCX: {e}")
            
    return statute_text, sample_text

# ==========================================
# Document Generation Logic (Word / DOCX) 
# Layout: 40% Left Margin, 60% Content
# ==========================================
def create_docx(content):
    doc = Document()
    
    # Half A4 (A5 Landscape) Page Setup (210mm x 148.5mm)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(148.5)
    
    # 40% of 210mm is 84mm. We set the left margin to 84mm to push text to the right 60%.
    section.left_margin = Mm(84) 
    section.right_margin = Mm(12)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    
    # Remaining usable width = 210 - 84 - 12 = 114mm

    # Set Default Font 
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Shruti' # Works natively with Gujarati in Word/Google Docs
    font.size = Pt(11)

    lines = content.split('\n')
    table_data = []
    in_table = False
    sig_buffer = []

    def flush_signatures():
        """Renders 3 side-by-side signatures using an invisible table fitting the 114mm width"""
        if sig_buffer:
            doc.add_paragraph().paragraph_format.space_before = Pt(20)
            
            sig_table = doc.add_table(rows=1, cols=3)
            sig_table.autofit = False
            # 114mm total width / 3 columns = 38mm each
            for cell in sig_table.columns[0].cells: cell.width = Mm(38)
            for cell in sig_table.columns[1].cells: cell.width = Mm(38)
            for cell in sig_table.columns[2].cells: cell.width = Mm(38)

            for i, sig in enumerate(sig_buffer):
                if i < 3:
                    p = sig_table.cell(0, i).paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    parts = sig.split(',')
                    for j, part in enumerate(parts):
                        run = p.add_run(part.strip())
                        if j < len(parts) - 1:
                            run.add_break()
                            
            sig_buffer.clear()

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        if line_stripped.startswith('|'):
            in_table = True
            row = [cell.strip() for cell in line_stripped.split('|') if cell.strip()]
            if not all(c == '-' for c in row[0].replace(' ', '')):
                table_data.append(row)
        else:
            if in_table:
                if table_data:
                    num_cols = len(table_data[0])
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_text
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            if row_idx == 0:
                                for run in p.runs:
                                    run.bold = True
                    
                    doc.add_paragraph()
                table_data = []
                in_table = False

            if line_stripped.startswith("તા.") or line_stripped.startswith("સ્થળ:"):
                flush_signatures()
                p = doc.add_paragraph(line_stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            elif "સાદર નોંધ" in line_stripped:
                flush_signatures()
                p = doc.add_paragraph()
                p.add_run(line_stripped).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            elif line_stripped.startswith("વિષય:"):
                flush_signatures()
                p = doc.add_paragraph()
                p.add_run(line_stripped).bold = True
            
            elif any(role in line_stripped for role in ["અધિકારી", "ઈન્ચાર્જ", "પ્રાધ્યાપક", "વડા"]) and not any(r in line_stripped for r in ["આચાર્ય", "ડીનશ્રી"]):
                sig_buffer.append(line_stripped)
            
            elif any(role in line_stripped for role in ["આચાર્ય", "ડીનશ્રી", "મહાવિધાયલય", "ન.કૃ.યુ"]):
                flush_signatures()
                
                doc.add_paragraph().paragraph_format.space_before = Pt(30)
                
                # Principal signature forced to the right of the 114mm content block
                p_table = doc.add_table(rows=1, cols=2)
                p_table.columns[0].width = Mm(40) # Spacer column
                p_table.columns[1].width = Mm(74)  # Signature column
                
                parts = line_stripped.split(",")
                formatted_line = "\n".join([p.strip() for p in parts])
                
                p = p_table.cell(0, 1).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(formatted_line)
            
            else:
                flush_signatures()
                doc.add_paragraph(line_stripped)

    flush_signatures()
    
    if in_table and table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# Streamlit App UI
# ==========================================
st.set_page_config(page_title="સાદર નોંધ જનરેટર", layout="wide")
st.title("સાદર નોંધ જનરેટર (Intelligent Sadar Nondh App)")

api_key = st.sidebar.text_input("Enter Gemini API Key", type="password")

tab1, tab2, tab3 = st.tabs(["નવી સાદર નોંધ બનાવો (Create New)", "જુની નોંધ (Archives)", "સેટિંગ્સ (Settings / Files)"])

with tab1:
    st.markdown("### જરૂરિયાતની વિગત આપો (Provide Requirements)")
    
    col1, col2 = st.columns(2)
    with col1:
        text_prompt = st.text_area("તમારી જરૂરિયાત લખો (Write short requirement):", 
                                   placeholder="e.g., need 10 entomological pins and 5 files for the AINP scheme.")
    with col2:
        uploaded_image = st.file_uploader("અથવા હાથથી લખેલી ચબરખીનો ફોટો અપલોડ કરો (Upload handwritten note):", type=["jpg", "jpeg", "png"])
    
    if st.button("જનરેટ કરો (Generate)"):
        if not api_key:
            st.error("Please enter your Gemini API Key in the sidebar.")
        elif not os.path.exists("121_Statutes_uploaded.pdf") or not os.path.exists("sample_nondh_uploaded.docx"):
             st.warning("Please upload the Statute PDF and Sample DOCX in the 'સેટિંગ્સ (Settings)' tab first.")
        elif not text_prompt and not uploaded_image:
            st.warning("Please provide either a text requirement or an image.")
        else:
            with st.spinner("સ્ટેચ્યુટ ૧૨૧ ની ચકાસણી અને નોંધ તૈયાર કરવામાં આવી રહી છે..."):
                try:
                    statute_context, sample_context = load_permanent_context()
                    
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel('gemini-3.1-flash-lite-preview')
                    
                    sys_prompt = f"""
                    You are an expert administrative AI for the Department of Entomology, N. M. College of Agriculture, NAU, Navsari.
                    Your task is to generate a formal 'સાદર નોંધ' in Gujarati based on the user's brief input or image. 
                    
                    CRITICAL INSTRUCTION: You MUST format the output EXACTLY matching the 'Sample Nondh Format' provided below.
                    
                    [CONTEXT START]
                    {statute_context[:15000]}
                    
                    {sample_context}
                    [CONTEXT END]

                    Format REQUIRED:
                    તા. {datetime.date.today().strftime('%d/%m/%Y')}
                    સ્થળ: નવસારી
                    સાદર નોંધ:
                    વિષય: [Appropriate Subject...]
                    સવિનય ઉપરોક્ત વિષય અન્વયે જણાવવાનું કે, અત્રેનાં કીટકશાસ્ત્ર વિભાગની આઈ.સી.એ.આર. યોજના AINP on Agril Acarology બ.સ. ૩૦૩/૨૦૯૨ અંતર્ગત [Detailed logical reason]. સદર વસ્તુનો કુલ અંદાજિત ખર્ચ [Total Amount] થનાર છે.
                    જે આપ સાહેબશ્રીને સ્ટેચ્યુટ ૧૨૧ની આઈટમ નંબર [Insert Correct Item No.] મુજબ એનાયત થયેલ સત્તા અનુસાર સૈદ્ધાંતિક મંજુરી આપવા વિનંતી. સદર ખર્ચ અત્રેના વિભાગમાં ચાલતી આઈ.સી.એ.આર યોજના (બ.સ. ૩૦૩/૨૦૯૨) માં કરવામાં આવશે.

                    [If multiple items, include a markdown table. Columns MUST be: ક્રમ | વિગત | જથ્થો | કિંમત | કુલ કિંમત]

                    ખેતીવાડી અધિકારી,કીટકશાસ્ત્ર વિભાગ
                    પ્રોજેકટ ઈન્ચાર્જ,કીટકશાસ્ત્ર વિભાગ
                    પ્રાધ્યાપક અને વડા,કીટકશાસ્ત્ર વિભાગ

                    આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિધાયલય, ન.કૃ.યુ. નવસારી
                    """
                    
                    inputs = [sys_prompt, text_prompt]
                    if uploaded_image:
                        img = Image.open(uploaded_image)
                        inputs.append(img)
                        
                    response = model.generate_content(inputs)
                    st.session_state['generated_nondh'] = response.text
                    st.success("સાદર નોંધ સફળતાપૂર્વક તૈયાર થઈ ગઈ છે!")
                    
                except Exception as e:
                    st.error(f"Error generating document: {e}")

    if 'generated_nondh' in st.session_state:
        st.markdown("---")
        st.markdown("### ડ્રાફ્ટ એડિટિંગ અને પ્રીવ્યુ (Draft Editing & Live Preview)")
        st.info("નોંધ: ઉપરના ખાનામાં ટેક્સ્ટ/માર્કડાઉન બદલો. નીચેના ભાગમાં તે આપોઆપ 40/60 લેઆઉટમાં ટેબલ સાથે દેખાશે.")
        
        # Text editor for raw markdown
        edited_text = st.text_area("અહીં ડ્રાફ્ટમાં સુધારા-વધારા કરો (Edit Raw Text / Table):", 
                                   st.session_state['generated_nondh'], height=250)
        
        # Live Visual Preview split into 40% space and 60% content to match the Word Doc
        st.markdown("#### દસ્તાવેજ પ્રીવ્યુ (Visual Preview)")
        with st.container(border=True):
            prev_col1, prev_col2 = st.columns([4, 6]) # 40% blank left, 60% content right
            with prev_col2:
                st.markdown(edited_text) # This natively renders the Markdown Table beautifully
        
        st.markdown("---")
        col_save, col_down = st.columns(2)
        with col_save:
            if st.button("આર્કાઇવમાં સેવ કરો (Save & Approve)"):
                subject_line = "No Subject"
                for line in edited_text.split('\n'):
                    if "વિષય:" in line:
                        subject_line = line.replace("વિષય:", "").strip()
                        break
                save_to_db(subject_line, edited_text)
                st.success("નોંધ ડેટાબેઝમાં સાચવી લેવામાં આવી છે!")
                
        with col_down:
            docx_data = create_docx(edited_text)
            st.download_button(label="Download as Word (DOCX)",
                               data=docx_data,
                               file_name=f"Sadar_Nondh_{datetime.date.today().strftime('%d_%m_%Y')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab2:
    st.markdown("### જુની નોંધ શોધો (Search Archives)")
    
    current_year = datetime.date.today().year
    years = [str(y) for y in range(current_year-2, current_year+3)]
    months = ["All", "01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
    
    col_y, col_m = st.columns(2)
    with col_y:
        sel_year = st.selectbox("વર્ષ (Year):", years, index=2)
    with col_m:
        sel_month = st.selectbox("મહિનો (Month):", months)
        
    if st.button("શોધો (Search)"):
        records = get_archives(sel_month, sel_year)
        if records:
            for idx, record in enumerate(records):
                date, subject, content = record
                with st.expander(f"{date} - {subject}"):
                    # Show preview in archive as well using the 40/60 split
                    arc_col1, arc_col2 = st.columns([4, 6])
                    with arc_col2:
                        st.markdown(content)
                    
                    archived_docx = create_docx(content)
                    st.download_button(label="Download this Document (Word)",
                                       data=archived_docx,
                                       file_name=f"Archive_{date.replace('/', '_')}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       key=f"dl_{idx}")
        else:
            st.info("કોઈ રેકોર્ડ મળેલ નથી (No records found for this period).")

with tab3:
    st.markdown("### કાયમી ફાઇલો અપલોડ કરો (Upload Permanent Reference Files)")
    st.info("અહીં એકવાર અપલોડ કરેલી ફાઇલો એપમાં સેવ થઈ જશે અને દર વખતે નવી નોંધ બનાવતી વખતે બેકગ્રાઉન્ડમાં ઉપયોગમાં લેવાશે.")
    
    col_pdf, col_docx = st.columns(2)
    
    with col_pdf:
        statute_file = st.file_uploader("સ્ટેચ્યુટ પીડીએફ (Statute 121 PDF):", type=["pdf"])
        if statute_file:
            if st.button("Save Statute PDF"):
                with open("121_Statutes_uploaded.pdf", "wb") as f:
                    f.write(statute_file.getbuffer())
                load_permanent_context.clear()
                st.success("Statute PDF saved permanently!")
        
        if os.path.exists("121_Statutes_uploaded.pdf"):
            st.success("✅ Statute PDF is currently saved and active.")

    with col_docx:
        sample_file = st.file_uploader("નમૂનાની વર્ડ ફાઇલ (Sample Nondh DOCX):", type=["docx"])
        if sample_file:
            if st.button("Save Sample DOCX"):
                with open("sample_nondh_uploaded.docx", "wb") as f:
                    f.write(sample_file.getbuffer())
                load_permanent_context.clear()
                st.success("Sample DOCX saved permanently!")
                
        if os.path.exists("sample_nondh_uploaded.docx"):
            st.success("✅ Sample DOCX is currently saved and active.")
